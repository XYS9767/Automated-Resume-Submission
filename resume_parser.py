"""
简历解析模块 — 支持 PDF / DOCX / TXT 格式
提取：姓名、联系方式、技能标签、工作经历、教育背景
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from utils import resolve_path, clean_text, get_logger

logger = get_logger("resume")


# ---------- 数据结构 ----------

@dataclass
class Experience:
    """单段工作/教育经历"""
    company_or_school: str = ""
    title_or_major: str = ""
    start_date: str = ""
    end_date: str = ""
    description: str = ""


@dataclass
class ResumeData:
    """解析后的简历数据结构"""
    name: str = ""
    phone: str = ""
    email: str = ""
    skills: list[str] = field(default_factory=list)
    work_experiences: list[Experience] = field(default_factory=list)
    education: list[Experience] = field(default_factory=list)
    raw_text: str = ""
    parsed_from: str = ""

    def to_dict(self) -> dict:
        return {
            "name": self.name,
            "phone": self.phone,
            "email": self.email,
            "skills": self.skills,
            "work_experiences": [
                {"company": e.company_or_school, "title": e.title_or_major,
                 "period": f"{e.start_date} - {e.end_date}",
                 "description": e.description}
                for e in self.work_experiences
            ],
            "education": [
                {"school": e.company_or_school, "major": e.title_or_major,
                 "period": f"{e.start_date} - {e.end_date}"}
                for e in self.education
            ],
        }


# ---------- 解析器 ----------

class ResumeParser:
    """简历解析器"""

    # 技能词库（英文 + 中文 + 框架 + 嵌入式）
    SKILL_KEYWORDS = [
        # 编程语言
        "Python", "Java", "Go", "Golang", "C++", "C#", "PHP", "Ruby",
        "JavaScript", "TypeScript", "Kotlin", "Swift", "Rust", "Scala",
        "Shell", "Bash", "C", "汇编", "Assembly", "Verilog", "VHDL", "MATLAB",
        # 前端
        "React", "Vue", "Angular", "HTML5", "CSS3", "jQuery", "Bootstrap",
        "Webpack", "Vite", "微信小程序", "小程序",
        # 后端框架
        "Django", "Flask", "FastAPI", "Spring", "Spring Boot", "Spring Cloud",
        "MyBatis", "Hibernate", "Express", "NestJS", "Gin", "Beego",
        # 数据库
        "MySQL", "PostgreSQL", "MongoDB", "Redis", "Oracle", "SQL Server",
        "Elasticsearch", "Memcached", "HBase", "Hive", "ClickHouse",
        # 中间件 & 消息队列
        "Kafka", "RabbitMQ", "RocketMQ", "Zookeeper", "Nginx", "Tomcat",
        "Dubbo", "Nacos", "Sentinel", "Gateway",
        # 运维 & 容器
        "Docker", "Kubernetes", "K8s", "Jenkins", "GitLab CI", "GitHub Actions",
        "Terraform", "Ansible", "Prometheus", "Grafana",
        "CentOS", "Ubuntu", "LNMP", "Harbor", "Loki", "ELK", "TiDB", "Ceph",
        "HAProxy", "Traefik", "Keepalived", "Supervisord", "Alertmanager",
        "PromQL", "Exporter", "ServiceMonitor", "Webhook", "Docker Compose",
        "Dockerfile", "SystemD", "Ingress", "Deployment", "Scheduler",
        "GitLab", "DevOps", "ACK", "ACR", "DTS", "X-Pack",
        "awk", "sed", "Playbook", "Requests", "Telegram",
        "python-telegram-bot", "jenkinsapi", "watchdog", "InnoDB", "utf8mb4",
        # 基础
        "Linux", "Git", "HTTP", "TCP/IP", "WebSocket", "RESTful", "GraphQL", "gRPC",
        # 大数据 & AI
        "Hadoop", "Spark", "Flink", "Storm", "TensorFlow", "PyTorch",
        "Scikit-learn", "Pandas", "NumPy", "Matplotlib", "Jupyter",
        "MapReduce", "YARN", "HDFS", "SparkCore", "SparkSQL",
        "SparkStreaming", "SparkOnYARN", "ZooKeeper",
        "Azkaban", "Sqoop", "Flume", "Ganglia",
        # 云服务
        "AWS", "Azure", "阿里云", "腾讯云", "华为云",
        # 软技能 & 领域
        "微服务", "分布式", "高并发", "多线程", "爬虫", "数据分析", "机器学习",
        "敏捷开发", "CI/CD",
        # 中文常见描述
        "数据库", "后端开发", "前端开发", "全栈", "运维", "测试", "自动化",
        # 云原生 / 运维（参考运维工程师简历）
        "云原生", "容器化", "混合云", "负载均衡", "高可用", "日志分析",
        "监控告警", "自动化运维", "运维开发", "镜像仓库", "企业微信",
        "Web集群", "中间件", "缓存穿透", "缓存雪崩", "数仓", "数据同步",
        "日志采集", "全链路", "动态存储", "存储池", "弹性伸缩", "故障排查",
        # === 嵌入式 ===
        "嵌入式", "单片机", "ARM", "STM32", "GD32", "ESP32", "AVR", "PIC",
        "FPGA", "DSP", "Zynq", "Xilinx", "Altera",
        "RTOS", "FreeRTOS", "uCOS", "ThreadX", "RT-Thread", "VxWorks",
        "Linux驱动", "Linux内核", "BSP", "设备树", "Bootloader", "U-Boot",
        "I2C", "SPI", "UART", "USART", "CAN", "CAN-FD", "RS232", "RS485",
        "USB", "PCIe", "Ethernet", "SDIO", "PWM", "ADC", "DAC", "GPIO",
        "ZigBee", "BLE", "蓝牙", "WiFi", "LoRa", "NB-IoT", "4G", "5G",
        "Keil", "IAR", "GCC", "Makefile", "CMake", "交叉编译", "烧录",
        "示波器", "万用表", "逻辑分析仪", "硬件调试", "原理图", "PCB",
        "传感器", "电机控制", "电源管理", "FOC", "编码器",
        "物联网", "IoT", "智能家居", "工业控制", "汽车电子",
    ]

    # 章节关键词（支持多种写法；长关键词优先，避免「教育」误匹配）
    SECTION_KEYWORDS = {
        "work": [
            "工作经历", "工作经验", "项目经验", "项目经历",
            "实习经历", "社会实践", "Work Experience",
            "WORK EXPERIENCE", "PROJECTS",
        ],
        "edu": [
            "教育背景", "教育经历", "学习经历", "学历",
            "Education", "EDUCATION",
        ],
        "skills_section": [
            "技能", "专业技能", "技术栈", "技术能力", "个人技能",
            "Skills", "SKILLS", "Tech Stack",
        ],
    }

    # 日期范围：2020.01-2022.06 / 2020/01-至今 / 2021.9-2024.6
    DATE_RANGE_PATTERN = re.compile(
        r"(\d{4}\s*[年./-]?\s*\d{1,2}\s*[月]?)\s*"
        r"[-~–—至到]\s*"
        r"(\d{4}\s*[年./-]?\s*\d{1,2}\s*[月]?|至今|现在|Now|now)"
    )

    # 日期后的「公司 + 职位」拆分
    COMPANY_TITLE_PATTERN = re.compile(
        r"^(.+?)\s+([\u4e00-\u9fffA-Za-z0-9]{2,24}"
        r"(?:工程师|经理|主管|专员|负责人|总监|开发|运维|架构师|顾问|分析师|技术))"
        r"(?:\s*[\u4e00-\u9fffA-Za-z0-9/]+)?$",
        re.IGNORECASE,
    )

    DEGREE_KEYWORDS = ["本科", "硕士", "博士", "大专", "研究生", "学士", "MBA", "高中", "专科"]

    def __init__(self, file_path: str):
        self.file_path = resolve_path(file_path)
        self.logger = logger

    def parse(self) -> ResumeData:
        """解析简历，返回结构化数据"""
        if not self.file_path.exists():
            raise FileNotFoundError(f"简历文件不存在: {self.file_path}")

        ext = self.file_path.suffix.lower()
        raw_text = self._read_file(ext)
        self.logger.info(f"原文长度: {len(raw_text)} 字符")

        resume = ResumeData(raw_text=raw_text, parsed_from=self.file_path.name)

        self._extract_contact(resume)
        self._extract_skills(resume)
        self._extract_sections(resume)

        # 如果结构化提取的工作经历为空，走全文本兜底
        if len(resume.work_experiences) == 0:
            self.logger.info("结构化提取工作经历为 0，尝试全局扫描...")
            self._parse_work_experiences(resume, raw_text.split("\n"))

        self.logger.info(
            f"解析完成: {resume.name or '未知姓名'} | "
            f"技能 {len(resume.skills)} 项 | "
            f"工作经历 {len(resume.work_experiences)} 段 | "
            f"教育 {len(resume.education)} 段"
        )
        return resume

    def _read_file(self, ext: str) -> str:
        """根据后缀读取文件内容，取提取质量最好的方式"""
        if ext == ".pdf":
            return self._read_pdf_best()
        elif ext == ".docx":
            return self._read_docx()
        elif ext == ".txt":
            return self._read_txt()
        else:
            raise ValueError(f"不支持的简历格式: {ext}")

    def _read_pdf_best(self) -> str:
        """PDF: 同时尝试 pdfplumber 和 PyMuPDF，返回文本更长的那个"""
        text_plumber = ""
        text_mupdf = ""

        # 方式 1: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(str(self.file_path)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
            text_plumber = "\n".join(pages)
            self.logger.debug(f"pdfplumber: {len(text_plumber)} 字符")
        except Exception as e:
            self.logger.debug(f"pdfplumber 失败: {e}")

        # 方式 2: PyMuPDF
        try:
            import fitz
            doc = fitz.open(str(self.file_path))
            text_mupdf = "\n".join(page.get_text() for page in doc)
            doc.close()
            self.logger.debug(f"PyMuPDF: {len(text_mupdf)} 字符")
        except Exception as e:
            self.logger.debug(f"PyMuPDF 失败: {e}")

        # 取文本量更大的
        if len(text_mupdf) > len(text_plumber):
            self.logger.info(f"使用 PyMuPDF 提取结果 ({len(text_mupdf)} 字符)")
            return text_mupdf
        elif len(text_plumber) > 0:
            self.logger.info(f"使用 pdfplumber 提取结果 ({len(text_plumber)} 字符)")
            return text_plumber
        else:
            self.logger.warning("PDF 文本提取为空！可能是图片型 PDF，建议转换为文本或使用附件中的 TXT/DOCX 版本")
            return ""

    def _read_docx(self) -> str:
        """读取 DOCX"""
        from docx import Document
        doc = Document(str(self.file_path))
        lines = [p.text for p in doc.paragraphs]
        # 也读表格中的内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    lines.append(cell.text)
        return "\n".join(lines)

    def _read_txt(self) -> str:
        """读取 TXT（自动检测编码）"""
        encodings = ["utf-8", "gbk", "gb2312", "gb18030", "utf-16"]
        for enc in encodings:
            try:
                with open(self.file_path, "r", encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        with open(self.file_path, "rb") as f:
            return f.read().decode("utf-8", errors="ignore")

    # ---------- 信息提取 ----------

    def _extract_contact(self, resume: ResumeData):
        """提取姓名、手机号、邮箱"""
        text = resume.raw_text

        # 手机号
        phone_match = re.search(r"1[3-9]\d{9}", text)
        if phone_match:
            resume.phone = phone_match.group()

        # 邮箱
        email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
        if email_match:
            resume.email = email_match.group()

        # 姓名：扫描前 20 行非空文本
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in lines[:20]:
            # 显式 "姓名: xxx" / "Name: xxx"
            m = re.search(r"(?:姓名|Name)[：:]\s*(\S{2,4})", line)
            if m:
                resume.name = m.group(1)
                return
        # 兜底：第 1~3 行中纯中文且不是标题词
        skip_words = ["简历", "resume", "cv", "个人简介", "求职", "应聘", "联系", "电话", "邮箱"]
        for line in lines[:5]:
            clean = line.strip()
            # 2-4 个中文字符、不含符号
            if 2 <= len(clean) <= 4 and re.fullmatch(r"[\u4e00-\u9fff·]+", clean):
                if not any(sk in clean.lower() for sk in skip_words):
                    resume.name = clean
                    return

    def _extract_skills(self, resume: ResumeData):
        """从全文匹配技能关键词"""
        text_lower = resume.raw_text.lower()
        matched = set()
        for skill in self.SKILL_KEYWORDS:
            if skill.lower() in text_lower:
                matched.add(skill)
        resume.skills = sorted(matched, key=lambda s: self.SKILL_KEYWORDS.index(s) if s in self.SKILL_KEYWORDS else 999)

    def _is_section_header(self, line: str) -> Optional[str]:
        """判断一行是否为章节标题，返回章节名或 None"""
        stripped = line.strip()
        if not stripped or len(stripped) > 20:
            return None
        for section, keywords in self.SECTION_KEYWORDS.items():
            if section == "skills_section":
                continue
            for kw in sorted(keywords, key=len, reverse=True):
                if stripped == kw or stripped.replace(" ", "") == kw:
                    return section
        return None

    def _detect_section(self, line: str) -> Optional[str]:
        """检测行内是否含章节关键词（长词优先）"""
        stripped = line.strip()
        header = self._is_section_header(stripped)
        if header:
            return header
        for section, keywords in self.SECTION_KEYWORDS.items():
            if section == "skills_section":
                continue
            for kw in sorted(keywords, key=len, reverse=True):
                if kw in stripped and len(stripped) <= len(kw) + 4:
                    return section
        return None

    def _extract_sections(self, resume: ResumeData):
        """按章节关键词分割文本"""
        lines = resume.raw_text.split("\n")

        current_section = None
        work_lines = []
        edu_lines = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            section_hit = self._detect_section(stripped)
            if section_hit:
                current_section = section_hit
                continue  # 章节标题行不参与解析

            if current_section == "work":
                work_lines.append(stripped)
            elif current_section == "edu":
                edu_lines.append(stripped)

        self._parse_work_experiences(resume, work_lines)
        self._parse_education(resume, edu_lines)

    def _extract_sections_simple(self, resume: ResumeData):
        """简单文本的章节提取（TXT 回退方案）"""
        self._extract_sections(resume)

    def _split_company_title(self, text: str) -> tuple[str, str]:
        """从日期后的文本拆分公司与职位"""
        text = clean_text(text)
        if not text:
            return "", ""

        m = self.COMPANY_TITLE_PATTERN.match(text)
        if m:
            return m.group(1).strip(), m.group(2).strip()

        for kw in ["工程师", "开发", "经理", "主管", "实习", "专员", "负责人", "总监", "运维"]:
            idx = text.rfind(kw)
            if idx > 0:
                return text[:idx].strip(), text[idx:].strip()

        return text, ""

    def _fill_experience_from_date_line(
        self, line: str, dm: re.Match, exp: Experience
    ):
        """根据含日期行填充公司/职位（支持日期在前或在后）"""
        before = line[:dm.start()].strip()
        after = line[dm.end():].strip()

        if before:
            company, title = self._split_company_title(before)
            exp.company_or_school = company or clean_text(before)
            if title:
                exp.title_or_major = title
        elif after:
            company, title = self._split_company_title(after)
            exp.company_or_school = company or clean_text(after)
            exp.title_or_major = title

    def _parse_work_experiences(self, resume: ResumeData, lines: list[str]):
        """从行列表中解析工作经历"""
        if not lines:
            return

        date_pattern = self.DATE_RANGE_PATTERN
        current = Experience()
        for line in lines:
            line = re.sub(r"^[\uf06c\u2022●•·\-\*]\s*", "", line.strip())
            if not line:
                continue

            dm = date_pattern.search(line)
            if dm:
                if current.start_date and (current.company_or_school or current.title_or_major):
                    resume.work_experiences.append(current)
                current = Experience(
                    start_date=dm.group(1).strip(),
                    end_date=dm.group(2).strip(),
                )
                self._fill_experience_from_date_line(line, dm, current)
            elif current.start_date:
                stripped = clean_text(line)
                if stripped:
                    current.description += stripped + " "

        if current.start_date and (current.company_or_school or current.title_or_major):
            resume.work_experiences.append(current)

        for exp in resume.work_experiences:
            if not exp.title_or_major:
                for keyword in ["工程师", "开发", "经理", "主管", "实习", "专员", "负责人", "技术总监"]:
                    if keyword in exp.description or keyword in exp.company_or_school:
                        exp.title_or_major = keyword
                        break
            if not exp.description:
                exp.description = ""

    def _parse_education(self, resume: ResumeData, lines: list[str]):
        """解析教育背景"""
        if not lines:
            return

        date_pattern = self.DATE_RANGE_PATTERN
        current = Experience()
        for line in lines:
            line = re.sub(r"^[\uf06c\u2022●•·\-\*]\s*", "", line.strip())
            if not line:
                continue

            dm = date_pattern.search(line)
            if dm:
                if current.start_date and current.company_or_school:
                    resume.education.append(current)
                current = Experience(
                    start_date=dm.group(1).strip(),
                    end_date=dm.group(2).strip(),
                )
                after = clean_text(line[dm.end():])
                before = clean_text(line[:dm.start()])
                body = after or before

                school, major = "", ""
                for dk in self.DEGREE_KEYWORDS:
                    if dk in body:
                        major = dk
                        break

                if major:
                    school = body.replace(major, "").strip(" ，,")
                else:
                    parts = body.rsplit(" ", 1)
                    if len(parts) == 2 and len(parts[1]) <= 12:
                        school, major = parts[0], parts[1]
                    else:
                        school = body

                current.company_or_school = school.strip(" ，,")
                current.title_or_major = major
            elif current.start_date:
                stripped = clean_text(line)
                if any(dk in stripped for dk in self.DEGREE_KEYWORDS) and not current.title_or_major:
                    current.title_or_major = stripped
                elif stripped:
                    current.description += stripped + " "

        if current.start_date and current.company_or_school:
            resume.education.append(current)


# ---------- 便捷入口 ----------

def parse_resume(file_path: str) -> ResumeData:
    """一键解析简历"""
    parser = ResumeParser(file_path)
    return parser.parse()
